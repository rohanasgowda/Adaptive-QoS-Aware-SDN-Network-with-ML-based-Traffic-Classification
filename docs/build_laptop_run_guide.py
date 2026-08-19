from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("AI_SDN_Laptop_Run_Guide.docx")

NAVY = "17324D"
BLUE = "246B8E"
TEAL = "17807A"
GOLD = "B7801C"
INK = "24313A"
MUTED = "5F6B73"
PALE_BLUE = "EAF3F7"
PALE_TEAL = "E8F5F2"
PALE_GOLD = "FFF5DA"
LIGHT = "F4F6F8"
WHITE = "FFFFFF"
BORDER = "CBD5DC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(8.7)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def shade_paragraph(paragraph, fill: str, border: str = BORDER) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), border)
        p_bdr.append(edge)
    p_pr.append(p_bdr)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    shade_paragraph(p, LIGHT)
    for index, line in enumerate(text.strip().splitlines()):
        run = p.add_run(line)
        set_run_font(run, "Consolas", 8.7, NAVY)
        if index < len(text.strip().splitlines()) - 1:
            run.add_break()


def add_note(doc: Document, label: str, text: str, fill=PALE_BLUE, accent=BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    shade_paragraph(p, fill, border=fill)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_step(doc: Document, title: str, detail: str | None = None) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_with_next = bool(detail)
    r = p.add_run(title)
    r.bold = True
    if detail:
        r = p.add_run(f" - {detail}")
        r.bold = False


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            p = row.cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)

    # Editorial-cover header pattern, adapted for a compact technical runbook.
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "AI-ASSISTED SDN ROUTING  |  LAPTOP RUN GUIDE"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.runs[0], size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("Project runbook  |  Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "PAGE")

    # Cover.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-ASSISTED SDN\nTRAFFIC ROUTING")
    set_run_font(r, size=30, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Laptop Setup and End-to-End Run Guide")
    set_run_font(r, size=16, color=BLUE, bold=True)

    add_note(
        doc,
        "Purpose",
        "Use this guide to install the project on a Windows laptop with WSL2, train the Random Forest model, run the Mininet topology, test all three controllers, and find the generated results.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Tested environment")
    set_run_font(r, size=10, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Windows + WSL2 Ubuntu 22.04 | Mininet 2.3.0 | Open vSwitch 2.17.9 | Python 3.10")
    set_run_font(r, size=10, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Project folder")
    set_run_font(r, size=10, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(r"C:\Users\ROHAN A S GOWDA\OneDrive\Dokumen\computer_network_rohan")
    set_run_font(r, "Consolas", 9, NAVY)

    add_page_break(doc)

    doc.add_heading("1. Before You Start", level=1)
    add_note(doc, "Recommended path", "Run the full simulation inside WSL2. Use Windows PowerShell only to install or open WSL; run Mininet and the controllers in Ubuntu terminals.")
    doc.add_heading("What you need", level=2)
    for item in (
        "Windows 10 version 2004 or newer, or Windows 11.",
        "Administrator access for the initial WSL installation.",
        "Internet access for Ubuntu packages and Python dependencies.",
        "At least 4 GB free disk space and 8 GB RAM recommended.",
        "The complete project folder on this laptop.",
    ):
        add_bullet(doc, item)

    doc.add_heading("How the run is organized", level=2)
    add_table(
        doc,
        ["Terminal", "Runs as", "Purpose"],
        [
            ["Ubuntu Terminal 1", "Normal WSL user", "Starts Ryu, POX, or the raw OpenFlow controller."],
            ["Ubuntu Terminal 2", "sudo for Mininet", "Starts the six-switch topology and traffic workload."],
            ["Windows PowerShell", "Normal/Admin", "Installs WSL and opens additional Ubuntu sessions."],
        ],
        [2050, 2050, 5260],
    )
    add_note(doc, "Important", "Only one controller can listen on TCP port 6633 at a time. Stop the current controller with Ctrl+C before starting the next one.", fill=PALE_GOLD, accent=GOLD)

    doc.add_heading("2. Install and Verify WSL2", level=1)
    add_step(doc, "Open PowerShell as Administrator.")
    add_step(doc, "Install Ubuntu 22.04 with WSL2.")
    add_code(doc, "wsl --install -d Ubuntu-22.04")
    add_step(doc, "Restart Windows if requested, then open Ubuntu and create your Linux username and password.")
    add_step(doc, "Verify that Ubuntu is using WSL version 2.")
    add_code(doc, "wsl -l -v")
    add_note(doc, "Expected", "The output should list Ubuntu-22.04 with VERSION 2. If it shows version 1, run: wsl --set-version Ubuntu-22.04 2")

    doc.add_heading("3. One-Time Ubuntu Setup", level=1)
    add_step(doc, "Open Ubuntu 22.04.")
    add_step(doc, "Update package information and install Mininet, Open vSwitch, iperf, Git, and Python tools.")
    add_code(doc, "sudo apt update\nsudo apt install -y mininet openvswitch-switch iperf python3-pip git")
    add_step(doc, "Start Open vSwitch and check that it responds.")
    add_code(doc, "sudo service openvswitch-switch start\nsudo ovs-vsctl show")
    add_step(doc, "Move to the project folder from inside WSL.")
    add_code(doc, 'cd "/mnt/c/Users/ROHAN A S GOWDA/OneDrive/Dokumen/computer_network_rohan"\npwd')
    add_note(doc, "Why the quotes matter", "The Windows username contains spaces. Keep the double quotes around the WSL project path.", fill=PALE_GOLD, accent=GOLD)
    add_step(doc, "Install the Python packages for the AI, plots, and Ryu controller.")
    add_code(doc, "python3 -m pip install --user -r requirements-mininet.txt\nexport PATH=\"$HOME/.local/bin:$PATH\"")
    add_step(doc, "Verify the main tools.")
    add_code(doc, "python3 --version\nmn --version\novs-vsctl --version | head -n 1\n$HOME/.local/bin/ryu-manager --version")

    doc.add_heading("Optional: make the PATH setting permanent", level=2)
    add_code(doc, 'echo \'export PATH="$HOME/.local/bin:$PATH"\' >> ~/.bashrc\nsource ~/.bashrc')

    doc.add_heading("4. Train the AI Model", level=1)
    add_step(doc, "Confirm that you are in the project root.")
    add_code(doc, 'cd "/mnt/c/Users/ROHAN A S GOWDA/OneDrive/Dokumen/computer_network_rohan"')
    add_step(doc, "Generate training data, train the Random Forest, and save the report.")
    add_code(doc, "python3 -m ai.train_model")
    add_note(doc, "Expected files", "models/traffic_rf.joblib, data/training_flows.csv, and results/model_report.json should exist after training.", fill=PALE_TEAL, accent=TEAL)

    doc.add_heading("5. Run a Quick Test Without Mininet", level=1)
    p = doc.add_paragraph("This test checks the AI classifier, weighted routing, congestion logic, controller comparison, tables, and graphs. It is the fastest way to confirm the Python side is healthy.")
    p.paragraph_format.space_after = Pt(8)
    add_code(doc, "python3 -m experiments.inspect_path_choices\npython3 -m experiments.run_comparison --events 240 --output-dir results")
    doc.add_heading("Expected comparison outputs", level=2)
    add_table(
        doc,
        ["File", "What it contains"],
        [
            ["comparison_summary.csv", "Controller-level throughput, delay, install time, CPU, reroutes, and congestion score."],
            ["comparison_raw.csv", "Per-flow measurements used to calculate the summary."],
            ["throughput_mbps.png", "Throughput comparison graph."],
            ["delay_ms.png", "Average delay comparison graph."],
            ["install_ms.png", "Routing-rule installation timing graph."],
            ["cpu_percent.png", "Controller CPU comparison graph."],
        ],
        [3100, 6260],
    )
    add_note(doc, "Checkpoint", "If this step succeeds, the Python dependencies and model are working. Continue to Mininet for the real virtual-network test.")

    doc.add_heading("6. Full Mininet Run: Ryu", level=1)
    doc.add_heading("Terminal 1 - start Ryu", level=2)
    add_code(doc, 'cd "/mnt/c/Users/ROHAN A S GOWDA/OneDrive/Dokumen/computer_network_rohan"\nexport PYTHONPATH="$PWD"\n$HOME/.local/bin/ryu-manager --ofp-tcp-listen-port 6633 controllers/ryu_ai_controller.py')
    add_note(doc, "Leave it running", "Do not close Terminal 1. Open a second Ubuntu terminal for Mininet.", fill=PALE_GOLD, accent=GOLD)
    doc.add_heading("Terminal 2 - run topology and traffic", level=2)
    add_code(doc, 'cd "/mnt/c/Users/ROHAN A S GOWDA/OneDrive/Dokumen/computer_network_rohan"\nsudo env PYTHONPATH="$PWD" python3 -m topologies.run_mininet_experiment \\\n  --controller-ip 127.0.0.1 \\\n  --controller-port 6633 \\\n  --profile mixed \\\n  --duration 30 \\\n  --output results/mininet_ryu_run.json')
    add_note(doc, "Success signal", "Mininet should report 0% ping loss and write results/mininet_ryu_run.json. The controller writes decisions to results/ryu_decisions.jsonl.", fill=PALE_TEAL, accent=TEAL)

    doc.add_heading("7. Full Mininet Run: POX", level=1)
    p = doc.add_paragraph("POX is installed as a separate source checkout because it is not part of the Python requirements file.")
    doc.add_heading("One-time POX setup", level=2)
    add_code(doc, "git clone --depth 1 https://github.com/noxrepo/pox.git ~/pox_tmp\ncp controllers/pox_ai_controller.py ~/pox_tmp/ext/pox_ai_controller.py")
    add_note(doc, "After code changes", "Copy controllers/pox_ai_controller.py into ~/pox_tmp/ext/ again before the next POX run so POX uses the latest controller code.")
    doc.add_heading("Terminal 1 - start POX", level=2)
    add_code(doc, 'cd "/mnt/c/Users/ROHAN A S GOWDA/OneDrive/Dokumen/computer_network_rohan"\nPYTHONPATH=".:$HOME/pox_tmp" python3 "$HOME/pox_tmp/pox.py" pox_ai_controller')
    doc.add_heading("Terminal 2 - run topology and traffic", level=2)
    add_code(doc, 'cd "/mnt/c/Users/ROHAN A S GOWDA/OneDrive/Dokumen/computer_network_rohan"\nsudo env PYTHONPATH="$PWD" python3 -m topologies.run_mininet_experiment \\\n  --controller-ip 127.0.0.1 \\\n  --controller-port 6633 \\\n  --profile mixed \\\n  --duration 30 \\\n  --output results/mininet_pox_run.json')
    add_note(doc, "Success signal", "Mininet should write results/mininet_pox_run.json. POX writes decisions to results/pox_decisions.jsonl.", fill=PALE_TEAL, accent=TEAL)

    doc.add_heading("8. Full Mininet Run: Raw Controller", level=1)
    doc.add_heading("Terminal 1 - start the raw OpenFlow controller", level=2)
    add_code(doc, 'cd "/mnt/c/Users/ROHAN A S GOWDA/OneDrive/Dokumen/computer_network_rohan"\npython3 -m controllers.raw_openflow_controller \\\n  --host 127.0.0.1 \\\n  --port 6633 \\\n  --metrics results/raw_decisions.jsonl')
    doc.add_heading("Terminal 2 - run topology and traffic", level=2)
    add_code(doc, 'cd "/mnt/c/Users/ROHAN A S GOWDA/OneDrive/Dokumen/computer_network_rohan"\nsudo env PYTHONPATH="$PWD" python3 -m topologies.run_mininet_experiment \\\n  --controller-ip 127.0.0.1 \\\n  --controller-port 6633 \\\n  --profile mixed \\\n  --duration 30 \\\n  --output results/mininet_raw_run.json')

    doc.add_heading("9. Repeat Runs and Traffic Profiles", level=1)
    p = doc.add_paragraph("The experiment accepts five profiles. Replace mixed in the Mininet command with one of these values:")
    add_table(
        doc,
        ["Profile", "Behavior under test"],
        [
            ["mixed", "Runs VoIP, video, file, and web workloads together."],
            ["voip", "Frequent small UDP packets; prioritizes low delay."],
            ["video", "Steady large UDP packets; prioritizes bandwidth and stable delay."],
            ["file", "Bulk TCP transfer; prioritizes maximum throughput."],
            ["web", "Irregular TCP bursts; balances delay and congestion."],
        ],
        [2000, 7360],
    )
    add_code(doc, "# Example: a 60-second file-transfer test\n--profile file --duration 60")
    add_note(doc, "Fair comparison", "Use the same profile and duration for Ryu, POX, and Raw. Give each run a different output filename so results are not overwritten.")

    doc.add_heading("10. Stop and Clean the Network", level=1)
    add_step(doc, "Wait for the Mininet experiment to finish.")
    add_step(doc, "In Terminal 1, press Ctrl+C to stop the controller.")
    add_step(doc, "Clean any leftover Mininet namespaces, links, and switches.")
    add_code(doc, "sudo mn -c")
    add_step(doc, "Confirm that no controller is still using TCP port 6633 before the next run.")
    add_code(doc, "sudo ss -ltnp | grep 6633 || echo \"Port 6633 is free\"")

    doc.add_heading("11. Find and Read the Results", level=1)
    add_code(doc, "ls -lh results\ncat results/mininet_ryu_run.json\ncat results/mininet_pox_run.json\ncat results/mininet_raw_run.json")
    add_table(
        doc,
        ["Metric", "Meaning"],
        [
            ["pingall_loss_percent", "Connectivity across all six hosts. A healthy run should normally be 0%."],
            ["pingall_seconds", "Time needed to establish and test host-to-host connectivity."],
            ["iperf_h1_h6", "Measured throughput between h1 and h6."],
            ["traffic_seconds", "Elapsed time for the selected workload."],
            ["*.jsonl decision log", "Per-flow classification, chosen path, rule timing, and reroute details."],
        ],
        [2700, 6660],
    )

    doc.add_heading("12. Troubleshooting", level=1)
    add_table(
        doc,
        ["Problem", "Fix"],
        [
            ["ryu-manager: command not found", "Use $HOME/.local/bin/ryu-manager, or export $HOME/.local/bin in PATH."],
            ["Address already in use on port 6633", "Stop the previous controller with Ctrl+C. Check with sudo ss -ltnp | grep 6633."],
            ["Mininet hosts cannot ping", "Confirm the controller is already running, then run sudo mn -c and retry."],
            ["Open vSwitch is not available", "Run sudo service openvswitch-switch start and sudo ovs-vsctl show."],
            ["ModuleNotFoundError for project modules", "cd to the project root and export PYTHONPATH=$PWD before launching the controller."],
            ["POX cannot find pox_ai_controller", "Copy the file into ~/pox_tmp/ext and launch the component name pox_ai_controller."],
            ["Permission denied from Mininet", "Run only the Mininet experiment and mn -c with sudo; run controllers as the normal WSL user."],
            ["Old or duplicate Mininet processes", "Stop controllers, run sudo mn -c, and start a fresh pair of terminals."],
            ["OneDrive path is difficult", "Keep the project path in double quotes. For faster Linux I/O, optionally copy the project into ~/computer_network_rohan."],
        ],
        [3000, 6360],
    )

    doc.add_heading("13. Final Success Checklist", level=1)
    for item in (
        "WSL reports Ubuntu-22.04 as version 2.",
        "Open vSwitch responds to sudo ovs-vsctl show.",
        "python3 -m ai.train_model creates the .joblib model.",
        "The synthetic comparison creates CSV, JSON, and PNG files.",
        "Ryu, POX, and Raw each connect to six OpenFlow switches.",
        "Each Mininet run writes its own JSON result file.",
        "pingall_loss_percent is 0.0 in a healthy run.",
        "sudo mn -c completes after testing.",
    ):
        add_bullet(doc, item)

    add_note(doc, "Known-good benchmark", "The verified WSL2 runs for this project reached 0% ping loss and approximately 19.0/18.9 Mbits/sec between h1 and h6 for all three controllers.", fill=PALE_TEAL, accent=TEAL)

    core = doc.core_properties
    core.title = "AI-Assisted SDN Traffic Routing - Laptop Setup and Run Guide"
    core.subject = "Step-by-step Windows WSL2, Mininet, Ryu, POX, and raw-controller runbook"
    core.author = "AI-Assisted SDN Routing Project"
    core.keywords = "SDN, Mininet, WSL2, Ryu, POX, OpenFlow, Random Forest"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
